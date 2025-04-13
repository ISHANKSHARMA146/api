from io import BytesIO
import logging
from PyPDF2 import PdfReader
from docx import Document
import pytesseract
from PIL import Image
import re
from zipfile import ZipFile
import xml.etree.ElementTree as ET
import tempfile
import os

logger = logging.getLogger(__name__)

def parse_file(file_buffer: BytesIO, filename: str) -> str:
    """
    Determines the file type (PDF, DOCX, DOC, or image) and extracts text accordingly.
    
    Args:
        file_buffer: File buffer of the uploaded file.
        filename: Name of the uploaded file.
        
    Returns:
        Extracted text content as a string.
    """
    try:
        if filename.lower().endswith(".pdf"):
            return parse_pdf(file_buffer)
        elif filename.lower().endswith(".docx"):
            return parse_docx(file_buffer)
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            return image_to_text(file_buffer)
        else:
            raise ValueError("Unsupported file format. Only PDF, DOCX, and image formats are supported.")
    except Exception as e:
        logger.error(f"Error parsing file '{filename}': {str(e)}", exc_info=True)
        raise

def parse_pdf(file_buffer: BytesIO) -> str:
    """
    Extracts text from a PDF file, including hyperlinks.
    
    Args:
        file_buffer: File buffer of the uploaded PDF file.
        
    Returns:
        Extracted text content as a string, including hyperlinks.
    """
    try:
        logger.info("Parsing PDF file")
        reader = PdfReader(file_buffer)
        text = ""
        hyperlinks = []

        # Extract text from each page and gather hyperlinks from metadata if available
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text

            # Extract hyperlinks from annotations (if available)
            if "/Annots" in page:
                annotations = page["/Annots"]
                for annotation in annotations:
                    # Check if annotation is an IndirectObject and resolve it properly
                    if isinstance(annotation, dict):
                        if "/A" in annotation and "/URI" in annotation["/A"]:
                            hyperlinks.append(annotation["/A"]["/URI"])

        # Join the hyperlinks into a single string (one per line)
        hyperlinks_text = '\n'.join(hyperlinks)
        return text.strip() + '\n' + hyperlinks_text

    except Exception as e:
        logger.error(f"Error reading PDF file: {str(e)}", exc_info=True)
        raise

def parse_docx(file_buffer: BytesIO) -> str:
    """
    Extracts text from a DOCX file, including hyperlinks and headers/footers.
    
    Args:
        file_buffer: File buffer of the uploaded DOCX file.
        
    Returns:
        Extracted text content as a string, including hyperlinks.
    """
    try:
        logger.info("Parsing DOCX file")
        doc = Document(file_buffer)
        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text += paragraph.text + '\n'

        # Extract hyperlinks if possible
        hyperlinks = []
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if hasattr(run, '_element') and run._element.xpath('./w:hyperlink'):
                    for hyperlink in run._element.xpath('./w:hyperlink'):
                        if 'r:id' in hyperlink.attrib:
                            hyperlinks.append(hyperlink.attrib['r:id'])

        # Extract header and footer text
        header_footer_text = ""
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                header_footer_text += paragraph.text + '\n'
            
            # Footer
            for paragraph in section.footer.paragraphs:
                header_footer_text += paragraph.text + '\n'

        return text.strip() + '\n' + '\n'.join(hyperlinks) + '\n' + header_footer_text.strip()

    except Exception as e:
        logger.error(f"Error reading DOCX file: {str(e)}", exc_info=True)
        raise

def image_to_text(file_buffer: BytesIO) -> str:
    """
    Extract text from an image using Tesseract OCR.
    
    Args:
        file_buffer: The image file buffer.
        
    Returns:
        Extracted text content as a string.
    """
    try:
        logger.info("Extracting text from image")
        image = Image.open(file_buffer)
        text = pytesseract.image_to_string(image)
        return text.strip()
    
    except Exception as e:
        logger.error(f"Error processing image file: {str(e)}", exc_info=True)
        raise 